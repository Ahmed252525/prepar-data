import os
import re
from docx import Document
from docx.oxml.ns import qn
import zipfile

INPUT_FOLDER = "word inmaa data"
OUTPUT_FOLDER = "markdown_split_pages"
os.makedirs(OUTPUT_FOLDER, exist_ok=True)


def word_to_markdown_pages(word_file, base_filename):
    """Convert Word document to multiple Markdown files split by page."""
    try:
        doc = Document(word_file)
    except Exception as e:
        raise RuntimeError(f"Cannot read file: {word_file} — {e}")

    markdown_pages = []
    current_lines = []
    page_number = 1

    current_lines.append(f"\n---\n**Page {page_number}**\n---\n")

    for element in doc.element.body:
        # Detect page break
        if any(child.tag == qn('w:br') and child.attrib.get(qn('w:type')) == 'page' for child in element.iter()):
            markdown_pages.append((page_number, current_lines))
            page_number += 1
            current_lines = [f"\n---\n**Page {page_number}**\n---\n"]

        if element.tag.endswith('p'):
            para = next((p for p in doc.paragraphs if p._element == element), None)
            if para and para.text.strip():
                text = para.text.strip()
                if para.style.name.startswith('Heading'):
                    level = para.style.name.split()[-1]
                    if level.isdigit():
                        current_lines.append('#' * int(level) + ' ' + text)
                    else:
                        current_lines.append('# ' + text)
                else:
                    current_lines.append(text)
                current_lines.append('')

        elif element.tag.endswith('tbl'):
            table = next((t for t in doc.tables if t._element == element), None)
            if table:
                markdown_table = process_table_proper_format(table)
                if markdown_table:
                    current_lines.extend(markdown_table)
                    current_lines.append('')

    if current_lines:
        markdown_pages.append((page_number, current_lines))

    for page_num, lines in markdown_pages:
        output_path = os.path.join(OUTPUT_FOLDER, f"{base_filename}_page_{page_num}.md")
        with open(output_path, 'w', encoding='utf-8') as f:
            f.write('\n'.join(lines))
        print(f"✅ Saved: {output_path}")


def is_valid_docx(path):
    try:
        with zipfile.ZipFile(path, 'r') as z:
            return 'word/document.xml' in z.namelist()
    except:
        return False


# ========== KEEP OLD FUNCTIONS =============

def process_table_proper_format(table):
    if not table.rows:
        return []
    result_lines = []
    has_nested_tables = any(cell.tables for row in table.rows for cell in row.cells)
    if has_nested_tables:
        result_lines = process_nested_table_content(table)
    else:
        result_lines = process_regular_table(table)
    return result_lines


def process_nested_table_content(table):
    all_tables = []
    for row in table.rows:
        for cell in row.cells:
            main_content = [para.text.strip() for para in cell.paragraphs if para.text.strip()]
            for nested_table in cell.tables:
                table_data = extract_table_data(nested_table)
                if table_data:
                    all_tables.append(table_data)

    result_lines = []
    for idx, table_data in enumerate(all_tables):
        if idx > 0:
            result_lines.append("")
        if table_data:
            header_row = table_data[0]
            result_lines.append('| ' + ' | '.join(header_row) + ' |')
            result_lines.append('| ' + ' | '.join(['---'] * len(header_row)) + ' |')
            for row in table_data[1:]:
                while len(row) < len(header_row):
                    row.append('–')
                row = row[:len(header_row)]
                result_lines.append('| ' + ' | '.join(row) + ' |')
    return result_lines


def process_regular_table(table):
    table_data = extract_table_data(table)
    if not table_data:
        return []
    result_lines = []
    header_row = table_data[0]
    result_lines.append('| ' + ' | '.join(header_row) + ' |')
    result_lines.append('| ' + ' | '.join(['---'] * len(header_row)) + ' |')
    for row in table_data[1:]:
        while len(row) < len(header_row):
            row.append('–')
        row = row[:len(header_row)]
        result_lines.append('| ' + ' | '.join(row) + ' |')
    return result_lines


def extract_table_data(table):
    if not table.rows:
        return []
    table_data = []
    for row in table.rows:
        row_data = []
        for cell in row.cells:
            cell_text_parts = [para.text.strip() for para in cell.paragraphs if para.text.strip()]
            cell_text = clean_cell_text(' '.join(cell_text_parts)) or '–'
            row_data.append(cell_text)
        if any(cell != '–' for cell in row_data):
            table_data.append(row_data)
    return table_data


def clean_cell_text(text):
    if not text:
        return ""
    text = text.replace('|', '&#124;').replace('\n', ' ').replace('\r', ' ').replace('\t', ' ')
    text = re.sub(r'\s+', ' ', text).strip()
    text = format_percentages(text)
    return text


def format_percentages(text):
    text = re.sub(r'(\d+)\s*%', r'\1%', text)
    if 'total' in text.lower():
        return 'Total'
    return text

# ========== BULK EXECUTION =============
if __name__ == "__main__":
    for filename in os.listdir(INPUT_FOLDER):
        if filename.endswith(".docx"):
            full_path = os.path.join(INPUT_FOLDER, filename)
            if not is_valid_docx(full_path):
                print(f"⚠️ Skipping invalid file: {filename}")
                continue
            base_name = os.path.splitext(filename)[0]
            print(f"\n🔄 Processing: {filename}")
            try:
                word_to_markdown_pages(full_path, base_name)
            except Exception as e:
                print(f"❌ Error in '{filename}': {str(e)}")
